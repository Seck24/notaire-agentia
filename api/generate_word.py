"""
generate_word.py - API FastAPI pour generer des documents Word notariaux
Lance avec: uvicorn generate_word:app --host 0.0.0.0 --port 8001
"""

import base64
import os
import re
import tempfile
from datetime import datetime

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


app = FastAPI(title="NotaireAgent Word Generator")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["POST"],
    allow_headers=["*"],
)

TITRES_ACTES = {
    "vente_immobiliere": "ACTE DE VENTE IMMOBILIERE",
    "constitution_societe": "ACTE DE CONSTITUTION DE SOCIETE",
    "succession": "ACTE DE NOTORIETE - DEVOLUTION SUCCESSORALE",
    "donation": "ACTE DE DONATION",
    "ouverture_credit": "ACTE D'OUVERTURE DE CREDIT",
}


class WordRequest(BaseModel):
    type_acte: str
    texte_acte: str
    cabinet_nom: str = "Etude Notariale"


def set_margins(doc, top=2.5, bottom=2.5, left=3.0, right=3.0):
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def add_header(doc, cabinet_nom):
    header = doc.sections[0].header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(cabinet_nom.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(107, 76, 42)
    run.bold = True
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line = p2.add_run("_" * 60)
    line.font.size = Pt(8)
    line.font.color.rgb = RGBColor(200, 168, 130)


def add_title(doc, type_acte):
    titre = TITRES_ACTES.get(type_acte, "ACTE NOTARIE")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(24)
    p.space_after = Pt(24)
    run = p.add_run(titre)
    run.font.name = "Times New Roman"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(42, 26, 16)


def clean_markdown(text):
    """Remove markdown bold/italic markers."""
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    return text


def markdown_to_word(doc, texte):
    """Convert markdown-ish text to formatted Word paragraphs."""
    for line in texte.split("\n"):
        stripped = line.strip()

        # Skip markdown separators
        if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped) or re.match(r'^_{3,}$', stripped):
            continue

        # Empty line = paragraph break
        if not stripped:
            doc.add_paragraph("")
            continue

        # Heading 2: ## Title
        if stripped.startswith("## "):
            heading_text = clean_markdown(stripped[3:]).strip().upper()
            p = doc.add_paragraph()
            p.space_before = Pt(18)
            p.space_after = Pt(8)
            run = p.add_run(heading_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(42, 26, 16)
            continue

        # Heading 3: ### Subtitle
        if stripped.startswith("### "):
            heading_text = clean_markdown(stripped[4:]).strip().upper()
            p = doc.add_paragraph()
            p.space_before = Pt(14)
            p.space_after = Pt(6)
            run = p.add_run(heading_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(42, 26, 16)
            continue

        # Heading 1: # Title (also handle)
        if stripped.startswith("# "):
            heading_text = clean_markdown(stripped[2:]).strip().upper()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.space_before = Pt(20)
            p.space_after = Pt(10)
            run = p.add_run(heading_text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(42, 26, 16)
            continue

        # Detect ALL CAPS lines as headings (notarial style)
        cleaned = clean_markdown(stripped)
        is_heading = (
            cleaned.isupper()
            and len(cleaned) < 80
            and not cleaned.startswith("[")
        )
        is_article = cleaned.startswith("Article ") or cleaned.startswith("ARTICLE ")

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(16)

        if is_heading or is_article:
            p.space_before = Pt(14)
            p.space_after = Pt(6)
            run = p.add_run(cleaned)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(42, 26, 16)
        else:
            run = p.add_run(cleaned)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(26, 26, 26)

            if "[A COMPLETER]" in cleaned or "[A REMPLIR]" in cleaned:
                run.font.color.rgb = RGBColor(200, 80, 20)
                run.font.highlight_color = 7


@app.post("/generate-word")
async def generate_word(payload: WordRequest):
    try:
        doc = Document()
        set_margins(doc)
        add_header(doc, payload.cabinet_nom)
        add_title(doc, payload.type_acte)
        markdown_to_word(doc, payload.texte_acte)

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)

        with open(tmp_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()

        os.unlink(tmp_path)

        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"acte_{payload.type_acte}_{date_str}.docx"

        return {"docx_base64": b64, "filename": filename}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health():
    return {"status": "ok", "service": "notariat-word-generator"}
